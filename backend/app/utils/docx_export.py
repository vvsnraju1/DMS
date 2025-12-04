"""
DOCX Export Utility
Converts HTML content from CKEditor to DOCX format using python-docx
"""
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from io import BytesIO
from typing import Optional


def html_to_docx(
    html_content: str,
    title: Optional[str] = None,
    doc_number: Optional[str] = None,
    department: Optional[str] = None
) -> BytesIO:
    """
    Convert HTML content to DOCX format
    
    Args:
        html_content: HTML string from CKEditor
        title: Document title (optional, added as header)
        doc_number: Document number (optional)
        department: Department name (optional)
    
    Returns:
        BytesIO buffer containing DOCX file
    """
    # Create Word document
    doc = DocxDocument()
    
    # Add header with metadata
    if title or doc_number:
        # Title
        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Doc number and department
        if doc_number or department:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = meta_para.add_run()
            if doc_number:
                run.add_text(f"Document No: {doc_number}")
            if department:
                run.add_text(f" | Department: {department}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add separator
        doc.add_paragraph()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process each element
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'blockquote', 'table']):
        if element.name.startswith('h'):
            # Headings
            level = int(element.name[1])
            heading = doc.add_heading(element.get_text(strip=True), level=level)
            
        elif element.name == 'p':
            # Paragraphs
            para = doc.add_paragraph()
            _add_formatted_text(para, element)
            
        elif element.name == 'blockquote':
            # Block quotes
            para = doc.add_paragraph(element.get_text(strip=True))
            para.style = 'Intense Quote'
            
        elif element.name == 'ul':
            # Unordered lists
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
                
        elif element.name == 'ol':
            # Ordered lists
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(strip=True), style='List Number')
                
        elif element.name == 'table':
            # Tables
            _add_table(doc, element)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def _add_formatted_text(paragraph, element):
    """Add formatted text from HTML element to Word paragraph"""
    # Handle inline formatting
    for content in element.descendants:
        if content.name == 'strong' or content.name == 'b':
            run = paragraph.add_run(content.get_text())
            run.bold = True
        elif content.name == 'em' or content.name == 'i':
            run = paragraph.add_run(content.get_text())
            run.italic = True
        elif content.name == 'u':
            run = paragraph.add_run(content.get_text())
            run.underline = True
        elif content.name == 'code':
            run = paragraph.add_run(content.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif content.name == 'a':
            # Links
            run = paragraph.add_run(content.get_text())
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        elif isinstance(content, str) and content.strip():
            # Plain text
            if not content.parent or content.parent.name in ['p', 'li', 'td', 'th']:
                paragraph.add_run(content)


def _add_table(doc, table_element):
    """Add HTML table to Word document"""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Count columns
    first_row = rows[0]
    cols = len(first_row.find_all(['th', 'td']))
    
    # Create Word table
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    
    # Fill cells
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell.get_text(strip=True)
                
                # Bold header cells
                if cell.name == 'th':
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def docx_to_html(docx_buffer: BytesIO) -> str:
    """
    Convert DOCX to HTML (basic conversion)
    
    Args:
        docx_buffer: BytesIO buffer containing DOCX file
    
    Returns:
        HTML string
    """
    doc = DocxDocument(docx_buffer)
    
    html_parts = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if heading
        if para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading ', '')
            if level.isdigit():
                html_parts.append(f'<h{level}>{text}</h{level}>')
            else:
                html_parts.append(f'<h1>{text}</h1>')
        else:
            # Check for bold/italic
            html = '<p>'
            for run in para.runs:
                content = run.text
                if run.bold and run.italic:
                    html += f'<strong><em>{content}</em></strong>'
                elif run.bold:
                    html += f'<strong>{content}</strong>'
                elif run.italic:
                    html += f'<em>{content}</em>'
                else:
                    html += content
            html += '</p>'
            html_parts.append(html)
    
    # Handle tables
    for table in doc.tables:
        table_html = '<table border="1"><tbody>'
        for row in table.rows:
            table_html += '<tr>'
            for cell in row.cells:
                table_html += f'<td>{cell.text}</td>'
            table_html += '</tr>'
        table_html += '</tbody></table>'
        html_parts.append(table_html)
    
    return '\n'.join(html_parts)

