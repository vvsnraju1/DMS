"""
DOCX generation from template blocks with token replacement
"""
import os
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re

from app.utils.template_tokens import replace_tokens, DEFAULT_TOKEN_VALUES


def html_to_docx_paragraph(doc: Document, html_content: str, token_values: Dict[str, str] = None, strict: bool = False):
    """
    Convert HTML content to DOCX paragraphs with A4 formatting
    
    Args:
        doc: python-docx Document object
        html_content: HTML string (may contain tokens)
        token_values: Dictionary of token values for replacement
        strict: If True, fail on missing tokens
    """
    if not html_content:
        return
    
    # Replace tokens if values provided
    if token_values:
        html_content = replace_tokens(html_content, token_values, strict)
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process each element
    for element in soup.children:
        if element.name == 'h1':
            p = doc.add_heading(element.get_text(), level=1)
            p.runs[0].font.size = Pt(18)
        elif element.name == 'h2':
            p = doc.add_heading(element.get_text(), level=2)
            p.runs[0].font.size = Pt(16)
        elif element.name == 'h3':
            p = doc.add_heading(element.get_text(), level=3)
            p.runs[0].font.size = Pt(14)
        elif element.name == 'h4':
            p = doc.add_heading(element.get_text(), level=4)
            p.runs[0].font.size = Pt(12)
        elif element.name == 'h5':
            p = doc.add_heading(element.get_text(), level=5)
            p.runs[0].font.size = Pt(12)
        elif element.name == 'h6':
            p = doc.add_heading(element.get_text(), level=6)
            p.runs[0].font.size = Pt(12)
        elif element.name == 'p':
            p = doc.add_paragraph()
            _add_text_with_formatting(p, element)
            # Set default font size for paragraphs
            if not p.runs:
                run = p.add_run()
                run.font.size = Pt(12)
            else:
                for run in p.runs:
                    if not run.font.size:
                        run.font.size = Pt(12)
        elif element.name == 'ul' or element.name == 'ol':
            _add_list(doc, element, element.name == 'ol')
        elif element.name == 'table':
            _add_table(doc, element)
        elif element.name == 'br':
            doc.add_paragraph()
        elif element.name:
            # Other block elements - convert to paragraph
            p = doc.add_paragraph()
            _add_text_with_formatting(p, element)
            # Set default font size
            for run in p.runs:
                if not run.font.size:
                    run.font.size = Pt(12)


def _add_text_with_formatting(paragraph, element):
    """Add text with formatting from HTML element"""
    for content in element.descendants:
        if isinstance(content, str) and content.strip():
            run = paragraph.add_run(content.strip())
        elif hasattr(content, 'name'):
            if content.name == 'strong' or content.name == 'b':
                run = paragraph.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em' or content.name == 'i':
                run = paragraph.add_run(content.get_text())
                run.italic = True
            elif content.name == 'u':
                run = paragraph.add_run(content.get_text())
                run.underline = True
            elif content.name == 'br':
                paragraph.add_run().add_break()
            else:
                paragraph.add_run(content.get_text())


def _add_list(doc: Document, list_element, ordered: bool = False):
    """Add list (ordered or unordered) to document"""
    for li in list_element.find_all('li', recursive=False):
        p = doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')
        _add_text_with_formatting(p, li)


def _add_table(doc: Document, table_element):
    """Add table to document - A4 sized, full width"""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Determine number of columns
    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    
    # Set table width to full page width (18cm for A4 with margins)
    total_width = Cm(22.0)
    col_width = total_width / max_cols
    
    for col_idx in range(max_cols):
        table.columns[col_idx].width = col_width
    
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                cell_obj = table.rows[i].cells[j]
                cell_obj.text = cell.get_text(strip=True)
                # Set font size
                for paragraph in cell_obj.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                # Bold for header cells
                if cell.name == 'th':
                    for paragraph in cell_obj.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def generate_docx_from_template(
    template_data: Dict,
    token_values: Dict[str, str] = None,
    output_path: str = None,
    strict: bool = False
) -> str:
    """
    Generate DOCX document from template blocks
    
    Args:
        template_data: Template data structure with metadata and blocks
        token_values: Dictionary of token values for replacement
        output_path: Path to save DOCX file (if None, generates path)
        strict: If True, fail on missing tokens
        
    Returns:
        Path to generated DOCX file
    """
    # Create new document
    doc = Document()
    
    # Set A4 page size and margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4 height
    section.page_width = Cm(21.0)   # A4 width
    section.left_margin = Cm(1.5)   # 15mm left margin
    section.right_margin = Cm(1.5)  # 15mm right margin
    section.top_margin = Cm(2.0)   # 20mm top margin
    section.bottom_margin = Cm(2.0) # 20mm bottom margin
    
    # Set document properties from metadata
    metadata = template_data.get('metadata', {})
    if metadata:
        doc.core_properties.title = metadata.get('template_title', 'Document')
        doc.core_properties.author = 'DMS System'
    
    # Process blocks in order
    blocks = template_data.get('blocks', [])
    sorted_blocks = sorted(blocks, key=lambda b: b.get('order', 0))
    
    for block in sorted_blocks:
        block_type = block.get('type', 'paragraph')
        html_content = block.get('html', '')
        
        if block_type == 'title':
            # Title page block - special formatting
            _add_title_page(doc, metadata, token_values, strict)
        elif block_type == 'heading':
            # Heading block
            heading_text = block.get('metadata', {}).get('heading_text', '')
            if heading_text:
                doc.add_heading(heading_text, level=1)
            html_to_docx_paragraph(doc, html_content, token_values, strict)
        elif block_type in ['paragraph', 'procedure', 'steps']:
            html_to_docx_paragraph(doc, html_content, token_values, strict)
        elif block_type == 'table':
            # Regular table
            html_to_docx_paragraph(doc, html_content, token_values, strict)
        elif block_type == 'signatory_table':
            # Signatory table - special handling
            _add_signatory_table(doc, metadata, token_values, strict)
        elif block_type == 'annexure_table':
            # Annexure table - special handling
            _add_annexure_table(doc, token_values, strict)
        elif block_type == 'image':
            # Image block
            image_path = block.get('metadata', {}).get('image_path')
            if image_path and os.path.exists(image_path):
                try:
                    doc.add_picture(image_path, width=Inches(5))
                except:
                    pass
        else:
            # Default: treat as paragraph
            html_to_docx_paragraph(doc, html_content, token_values, strict)
    
    # Generate output path if not provided
    if not output_path:
        template_code = metadata.get('template_code', 'template')
        revision = metadata.get('revision', '01')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{template_code}_rev{revision}_{timestamp}.docx"
        output_path = os.path.join("storage/templates/generated", filename)
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Save document
    doc.save(output_path)
    
    return output_path


def _add_title_page(doc: Document, metadata: Dict, token_values: Dict = None, strict: bool = False):
    """Add title page with template metadata - Full width like signatory table"""
    # Create main table with 4 columns
    main_table = doc.add_table(rows=4, cols=4)
    main_table.style = 'Light Grid Accent 1'
    main_table.autofit = False
    
    # Set table width to full page width (18cm) - same as signatory table
    total_width = Cm(18.0)
    # Column widths: 20% for logo, 26.67% each for other 3 columns
    main_table.columns[0].width = Cm(3.6)   # 20%
    main_table.columns[1].width = Cm(4.8)    # 26.67%
    main_table.columns[2].width = Cm(4.8)    # 26.67%
    main_table.columns[3].width = Cm(4.8)    # 26.67%
    
    # Top row: zerokost logo (rowspan 2) and SOP Title
    logo_cell = main_table.rows[0].cells[0]
    logo_cell.merge(main_table.rows[1].cells[0])  # Merge with row below
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add logo image
    logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'static', 'zerokost-logo.png')
    # Fallback to storage if not in static
    if not os.path.exists(logo_path):
        logo_path = os.path.join('storage', 'static', 'zerokost-logo.png')
    
    if os.path.exists(logo_path):
        try:
            # Add image to paragraph
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Cm(3.5))  # Adjust size to fit (3.5cm width)
        except Exception as e:
            # Fallback to text if image fails
            logo_run = logo_para.add_run('zerokost\nHealthcare Pvt. Ltd.')
            logo_run.font.size = Pt(12)
            logo_run.bold = True
    else:
        # Fallback to text if image not found
        logo_run = logo_para.add_run('zerokost\nHealthcare Pvt. Ltd.')
        logo_run.font.size = Pt(12)
        logo_run.bold = True
    
    # SOP Title cell (spans 3 columns)
    title_cell = main_table.rows[0].cells[1]
    title_cell.merge(main_table.rows[0].cells[2])
    title_cell.merge(main_table.rows[0].cells[3])
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('STANDARD OPERATING PROCEDURE')
    title_run.font.size = Pt(16)
    title_run.bold = True
    
    # Second row: Empty cells (logo continues)
    for i in range(1, 4):
        main_table.rows[1].cells[i].text = ''
    
    # Third row: SOP Title
    sop_title_label = main_table.rows[2].cells[0]
    sop_title_label.text = 'SOP Title'
    sop_title_label.paragraphs[0].runs[0].bold = True
    sop_title_label.paragraphs[0].runs[0].font.size = Pt(11)
    
    sop_title_value = main_table.rows[2].cells[1]
    sop_title_value.merge(main_table.rows[2].cells[2])
    sop_title_value.merge(main_table.rows[2].cells[3])
    if token_values:
        sop_title_value.text = replace_tokens(metadata.get('template_title', '{{TEMPLATE_TITLE}}'), token_values, strict)
    else:
        sop_title_value.text = metadata.get('template_title', '{{TEMPLATE_TITLE}}')
    sop_title_value.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Fourth row: SOP Number, CC Number
    sop_num_label = main_table.rows[3].cells[0]
    sop_num_label.text = 'SOP Number'
    sop_num_label.paragraphs[0].runs[0].bold = True
    sop_num_label.paragraphs[0].runs[0].font.size = Pt(11)
    
    sop_num_value = main_table.rows[3].cells[1]
    if token_values:
        sop_num_value.text = replace_tokens(metadata.get('template_code', '{{TEMPLATE_CODE}}'), token_values, strict)
    else:
        sop_num_value.text = metadata.get('template_code', '{{TEMPLATE_CODE}}')
    sop_num_value.paragraphs[0].runs[0].font.size = Pt(11)
    # Add gray background (simulated with shading)
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
    sop_num_value._element.get_or_add_tcPr().append(shading_elm)
    
    cc_label = main_table.rows[3].cells[2]
    cc_label.text = 'CC Number'
    cc_label.paragraphs[0].runs[0].bold = True
    cc_label.paragraphs[0].runs[0].font.size = Pt(11)
    
    cc_value = main_table.rows[3].cells[3]
    if token_values:
        cc_value.text = replace_tokens(metadata.get('cc_number', '{{CC_NUMBER}}'), token_values, strict)
    else:
        cc_value.text = metadata.get('cc_number', '{{CC_NUMBER}}')
    cc_value.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Add another row for Effective Date and Next Review Date
    date_row = main_table.add_row()
    eff_date_label = date_row.cells[0]
    eff_date_label.text = 'Effective Date'
    eff_date_label.paragraphs[0].runs[0].bold = True
    eff_date_label.paragraphs[0].runs[0].font.size = Pt(11)
    
    eff_date_value = date_row.cells[1]
    if token_values:
        eff_date_value.text = replace_tokens(metadata.get('effective_date', '{{EFFECTIVE_DATE}}'), token_values, strict)
    else:
        eff_date_value.text = metadata.get('effective_date', '{{EFFECTIVE_DATE}}')
    eff_date_value.paragraphs[0].runs[0].font.size = Pt(11)
    
    review_date_label = date_row.cells[2]
    review_date_label.text = 'Next Review Date'
    review_date_label.paragraphs[0].runs[0].bold = True
    review_date_label.paragraphs[0].runs[0].font.size = Pt(11)
    
    review_date_value = date_row.cells[3]
    if token_values:
        review_date_value.text = replace_tokens(metadata.get('next_review_date', '{{NEXT_REVIEW_DATE}}'), token_values, strict)
    else:
        review_date_value.text = metadata.get('next_review_date', '{{NEXT_REVIEW_DATE}}')
    review_date_value.paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_page_break()


def _add_signatory_table(doc: Document, metadata: Dict, token_values: Dict = None, strict: bool = False):
    """Add signatory table - Resized to fit page width with smaller fonts"""
    doc.add_heading('Document Signatory Page', level=1)
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    
    # Set table width to full page width (18cm) with fixed column widths
    total_width = Cm(18.0)
    # Column widths: 15%, 20%, 20%, 20%, 25% (resized to fit)
    col_widths = [Cm(2.7), Cm(3.6), Cm(3.6), Cm(3.6), Cm(4.5)]
    for col_idx, width in enumerate(col_widths):
        table.columns[col_idx].width = width
    
    # Header row - smaller font
    headers = ['Signature Type', 'Name', 'Designation', 'Department', 'Date & Time']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)  # Reduced from 11pt
    
    # Signatory rows - smaller fonts for tokens
    signatories = [
        ('Prepared By', 'SIGNATORY_PREPARED'),
        ('Checked By', 'SIGNATORY_CHECKED'),
        ('Approved By', 'SIGNATORY_APPROVED'),
    ]
    
    for i, (label, prefix) in enumerate(signatories, 1):
        # Signature Type cell - 9pt
        label_cell = table.rows[i].cells[0]
        label_cell.text = label
        for paragraph in label_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
        
        # Token cells - 8pt (smaller for tokens)
        if token_values:
            name_cell = table.rows[i].cells[1]
            name_cell.text = token_values.get(f'{prefix}_NAME', f'{{{{{prefix}_NAME}}}}')
            for paragraph in name_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            
            designation_cell = table.rows[i].cells[2]
            designation_cell.text = token_values.get(f'{prefix}_DESIGNATION', f'{{{{{prefix}_DESIGNATION}}}}')
            for paragraph in designation_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            
            dept_cell = table.rows[i].cells[3]
            dept_cell.text = token_values.get(f'{prefix}_DEPARTMENT', f'{{{{{prefix}_DEPARTMENT}}}}')
            for paragraph in dept_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            
            date_cell = table.rows[i].cells[4]
            date_cell.text = token_values.get(f'{prefix}_DATE', f'{{{{{prefix}_DATE}}}}')
            for paragraph in date_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
        else:
            # Placeholder tokens - 8pt
            tokens = [
                f'{{{{{prefix}_NAME}}}}',
                f'{{{{{prefix}_DESIGNATION}}}}',
                f'{{{{{prefix}_DEPARTMENT}}}}',
                f'{{{{{prefix}_DATE}}}}'
            ]
            for j, token in enumerate(tokens, 1):
                cell = table.rows[i].cells[j]
                cell.text = token
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)


def _add_annexure_table(doc: Document, token_values: Dict = None, strict: bool = False):
    """Add annexure table - A4 sized, full width"""
    doc.add_heading('List of Annexures', level=1)
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    
    # Set table width to full page width (18cm)
    total_width = Cm(18.0)
    table.columns[0].width = Cm(5.4)  # 30%
    table.columns[1].width = Cm(12.6)  # 70%
    
    # Header row
    header_cell_1 = table.rows[0].cells[0]
    header_cell_1.text = 'Annexure No.'
    header_cell_1.paragraphs[0].runs[0].bold = True
    header_cell_1.paragraphs[0].runs[0].font.size = Pt(11)
    
    header_cell_2 = table.rows[0].cells[1]
    header_cell_2.text = 'Title'
    header_cell_2.paragraphs[0].runs[0].bold = True
    header_cell_2.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Add empty rows (can be populated from token_values if provided)
    for i in range(3):
        row = table.add_row()
        row.cells[0].text = ''
        row.cells[1].text = ''
        # Set font size
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

